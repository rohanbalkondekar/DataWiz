import docx2txt
import io
import re
import os
import fitz
import json
import faiss
import pickle
import PyPDF2
import tabula
import easyocr
import requests
import numpy as np
import pandas as pd
from PIL import Image
import aspose.words as aw
from docx import Document
from bs4 import BeautifulSoup
from langchain.vectorstores import FAISS
from sentence_transformers import SentenceTransformer
from youtube_transcript_api import YouTubeTranscriptApi
from langchain.embeddings import HuggingFaceInstructEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter

folder_path = './data'
model = SentenceTransformer('all-MiniLM-L6-v2')
# embeddings = HuggingFaceInstructEmbeddings(model_name="hkunlp/instructor-xl")

ingested_data_directory = "./ingested_data"
if not os.path.exists(ingested_data_directory):
    os.makedirs(ingested_data_directory)

# ─── Helper Functions ─────────────────────────────────────────────────────────

textSplit = RecursiveCharacterTextSplitter(chunk_size=550,
                                           chunk_overlap=100,
                                           length_function=len)


def ingestText(text):
    # ----------- Saving as plain text
    directory = "./ingested_data/text"
    if not os.path.exists(directory):
        os.makedirs(directory)

    # Check if the file exists
    file_path = os.path.join(directory, "data.txt")
    if os.path.exists(file_path):
        try:
            # Read the existing data from the file
            with open(file_path, "r", encoding="utf-8") as file:
                existing_data = file.read()
        except IOError:
            print(f"Error reading from file: {file_path}")
            existing_data = ""

    # Merge the new data with the existing data
        merged_data = existing_data + text

    # Write the merged data back to the file
        with open(file_path, "w", encoding="utf-8") as file:
            file.write(merged_data)
    else:
        # Create a new file and write the data
        with open(file_path, "w", encoding="utf-8") as file:
            file.write(text)

    # ----- Splitting into chunks
    doc_list = textSplit.split_text(text)
    data = doc_list
    # list to store to retrive later
    new_list = data

    # Directory to store the list
    directory = "./ingested_data/list"
    if not os.path.exists(directory):
        os.makedirs(directory)

    # File path for storing the list
    file_path = os.path.join(directory, "list.pkl")

    if os.path.exists(file_path):
        # Load the existing list from the pickle file
        with open(file_path, "rb") as file:
            existing_list = pickle.load(file)

    # Extend the existing list with the new list
        existing_list.extend(new_list)

    # Store the extended list as a pickle file
        with open(file_path, "wb") as file:
            pickle.dump(existing_list, file)
    else:
        # Store the new list as a pickle file
        with open(file_path, "wb") as file:
            pickle.dump(new_list, file)

    encoded_data = model.encode(data)
    index = faiss.IndexIDMap(faiss.IndexFlatIP(384))
    index.add_with_ids(encoded_data, np.array(range(0, len(data))))
    # vectorstore = FAISS.from_texts(texts=data, embedding=embeddings)

    # --------- Saving as FAISS index
    directory = "./ingested_data/index"
    if not os.path.exists(directory):
        os.makedirs(directory)

    faiss.write_index(index, './ingested_data/index/index')

# ─── Handle Dataframe ─────────────────────────────────────────────────────────


def ingestDataframe(df):
    # Create a directory to store the DataFrame
    directory = "./ingested_data/dataframe"
    if not os.path.exists(directory):
        os.makedirs(directory)

    # Define the file path for storing the DataFrame
    file_path = os.path.join(directory, "data.pkl")

    if os.path.exists(file_path):
        # Load the existing DataFrame from the pickle file
        with open(file_path, "rb") as file:
            existing_df = pickle.load(file)

    # Concatenate the existing DataFrame and the new DataFrame
        merged_df = pd.concat([existing_df, df], ignore_index=True)

    # Store the merged DataFrame as a pickle file
        with open(file_path, "wb") as file:
            pickle.dump(merged_df, file)
    else:
        # Store the DataFrame as a pickle file
        with open(file_path, "wb") as file:
            pickle.dump(df, file)

    df_text = df.to_string()

    # Removing excessive blank spaces
    df_text = re.sub(r'\s{10,}', ' ', df_text)
    df_text = re.sub(r'\n{8,}', '\n', df_text)

    df_json = df.to_json(orient='records')

    text = df_text + df_json

    ingestText(text)

# ─── For Excel Files ──────────────────────────────────────────────────────────


def ingestExcel(folder_path):
    def read_excel_files(folder_path):
        file_list = os.listdir(folder_path)
        df = pd.DataFrame()
        for file_name in file_list:
            if file_name.endswith('.xlsx') or file_name.endswith('.xls'):
                file_path = os.path.join(folder_path, file_name)
                excel_df = pd.read_excel(file_path)
                df = df._append(excel_df, ignore_index=True)
        df = df.fillna('')
        return df

    df = read_excel_files(folder_path)

    ingestDataframe(df)


# ─── For Text Files ───────────────────────────────────────────────────────────

def ingestTextFiles(folder_path):
    def read_txt_files(folder_path):
        file_list = os.listdir(folder_path)
        text = ""
        for file_name in file_list:
            if file_name.endswith('.txt'):
                file_path = os.path.join(folder_path, file_name)
                with open(file_path, "r", encoding="utf-8") as file:
                    file_text = file.read()
                    text += file_text + "\n"
        return text

    text = read_txt_files(folder_path)

    ingestText(text)


# ─── For Pdf Files ────────────────────────────────────────────────────────────

def ingestPDF(folder_path):
    def read_pdf_files(folder_path):
        file_list = os.listdir(folder_path)
        text = ""
        for file_name in file_list:
            if file_name.endswith('.pdf'):
                file_path = os.path.join(folder_path, file_name)
                with open(file_path, "rb") as file:
                    pdf_reader = PyPDF2.PdfFileReader(file)
                    for page in range(pdf_reader.numPages):
                        page_obj = pdf_reader.getPage(page)
                        page_text = page_obj.extractText()
                        text += page_text + "\n"
        return text

    # Read text from PDF files
    pdf_text_data = read_pdf_files(folder_path)

    ingestText(pdf_text_data)

    # Extract Tables from PDF files
    def extract_tables(folder_path):
        df = pd.DataFrame()
        for file_name in os.listdir(folder_path):
            if file_name.endswith(".pdf"):
                file_path = os.path.join(folder_path, file_name)

                try:
                    tables = tabula.read_pdf(
                        file_path, pages="all", multiple_tables=True)
                    for table in tables:
                        # Convert each table to a DataFrame
                        excel_df = pd.DataFrame(table)
                        # Append to the main DataFrame
                        df = pd.concat([df, excel_df], ignore_index=True)
                except Exception as e:
                    print(f"Error processing file {file_name}: {e}")

        df = df.fillna('')
        return df

    df = extract_tables(folder_path)
    ingestDataframe(df)


# ─── For Docx Files ───────────────────────────────────────────────────────────
def ingestDocx(folder_path):
    def read_docx_files(folder_path):
        file_list = os.listdir(folder_path)
        text = ""
        for file_name in file_list:
            if file_name.endswith('.docx'):
                file_path = os.path.join(folder_path, file_name)
                extracted_text = docx2txt.process(file_path)
                text += extracted_text + "\n"
        return text

    # Read text from DOCX files
    docx_data = read_docx_files(folder_path)
    ingestText(docx_data)

    def extract_tables_from_docx(file_path):
        docx_files = [file for file in os.listdir(
            folder_path) if file.endswith(".docx")]
        df = pd.DataFrame()

        for file in docx_files:
            file_path = os.path.join(folder_path, file)
            try:
                doc = Document(file_path)
                for table in doc.tables:
                    rows = []
                    for row in table.rows:
                        cells = [cell.text for cell in row.cells]
                        rows.append(cells)
                    docx_df = pd.DataFrame(rows)
                    df = pd.concat([df, docx_df], ignore_index=True)
            except Exception as e:
                print(f"Error extracting tables from {file}: {e}")
        df.fillna('', inplace=True)
        return df

    df = extract_tables_from_docx(folder_path)
    ingestDataframe(df)


# ─── For CSV Files ────────────────────────────────────────────────────────────
def ingestCSV(folder_path):
    def read_csv_files(folder_path):
        file_list = os.listdir(folder_path)
        df = pd.DataFrame()
        for file_name in file_list:
            if file_name.endswith('.csv'):
                file_path = os.path.join(folder_path, file_name)
                csv_df = pd.read_csv(file_path)
                df = df._append(csv_df, ignore_index=True)
        df = df.fillna('')
        return df

    df = read_csv_files(folder_path)
    ingestDataframe(df)


# ─── For Webpage ──────────────────────────────────────────────────────────

def ingestWebpage(links):

    def retrieve_text(link):
        response = requests.get(link)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, "html.parser")
            text = soup.get_text().strip()
            return text
        else:
            print(
                f"Error: Could not fetch the web page. Status code: {response.status_code}")
            return ""

    def scrape_links(links):
        txt_data = ""
        for link in links:
            txt_data += retrieve_text(link) + "\n\n"

            # Find all the anchor tags in the current page
            response = requests.get(link)
            soup = BeautifulSoup(response.content, "html.parser")
            anchor_tags = soup.find_all('a')

            # Extract the href attribute from each anchor tag
            nested_links = [tag['href']
                            for tag in anchor_tags if 'href' in tag.attrs]

            # Remove any duplicates and filter out non-http(s) links
            nested_links = list(set(nested_links))
            nested_links = [
                nested_link for nested_link in nested_links if nested_link.startswith('http')]

            # Retrieve text from nested links (first layer only)
            for nested_link in nested_links:
                nested_text = retrieve_text(nested_link)
                if nested_text:
                    txt_data += nested_text + "\n\n"

        return txt_data

    txt_data = scrape_links(links)
    # Clean up the text data
    txt_data = re.sub("\n{4,}", "\n", txt_data)

    ingestText(txt_data)


# ─── For Youtube Videos ───────────────────────────────────────────────────────
api_key = " " #use your youtube api key here


def ingest_youtube_video(links):
    transcription_str = ""
    for link in links:
        video_id = extract_video_id(link)
        if video_id is not None:
            transcription_str += get_video_info(video_id) + "\n\n"
            transcription = get_transcription(video_id)
            if transcription is not None:
                transcription_str += "Video Transcription: " + transcription + "\n\n\n"

    ingestText(transcription_str)


def extract_video_id(link):
    # Extracts the video ID from a YouTube link
    video_id = None
    if "youtube.com" in link:
        video_id = link.split("?v=")[1].split("&")[0]
    elif "youtu.be" in link:
        video_id = link.split("/")[-1]
    return video_id


def get_transcription(video_id):
    # Gets the transcription of a YouTube video
    try:
        transcription = YouTubeTranscriptApi.get_transcript(video_id)
        script = ""
        for text in transcription:
            t = text["text"]
            if t != '[Music]':
                script += t + " "
        return script
    except:
        return None


def get_video_info(video_id):
    url = f"https://www.googleapis.com/youtube/v3/videos?part=snippet&id={video_id}&key={api_key}"
    response = requests.get(url)
    if response.status_code == 200:
        data = json.loads(response.content.decode("utf-8"))
        if "items" in data and len(data["items"]) > 0:
            snippet = data["items"][0]["snippet"]
            video_info = f"Video ID: {video_id}\n"
            video_info += f"Title: {snippet['title']}\n"
            video_info += f"Description: {snippet['description']}\n"
            return video_info
    return None

# ─── Extract Images From Pdf And Docx And Perform Ocr ─────────────────────────


def extract_images_from_pdf(file_path, output_dir):
    # Create the output directory if it does not exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    pdf_file = fitz.open(file_path)
    extracted_images = []

    for page_index in range(len(pdf_file)):
        page = pdf_file[page_index]
        image_list = page.get_images(full=True)

        if image_list:
            print(
                f"[+] Found a total of {len(image_list)} images in page {page_index}")
        else:
            print(f"[!] No images found on page {page_index}")

        for image_index, img in enumerate(image_list, start=1):
            xref = img[0]
            base_image = pdf_file.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image = Image.open(io.BytesIO(image_bytes))

            min_width = 100
            min_height = 100
            output_format = "png"

            if image.width >= min_width and image.height >= min_height:
                image_path = os.path.join(
                    output_dir, f"image{page_index + 1}_{image_index}.{output_format}")
                image.save(open(image_path, "wb"),
                           format=output_format.upper())
                extracted_images.append(image_path)
            else:
                print(
                    f"[-] Skipping image {image_index} on page {page_index} due to its small size.")

    return extracted_images


def nextract_images_from_docx(file_path, output_dir):
    doc = aw.Document(file_path)
    shapes = doc.get_child_nodes(aw.NodeType.SHAPE, True)
    image_index = 0
    extracted_images = []

    for shape in shapes:
        shape = shape.as_shape()
        if shape.has_image:
            image_file_name = f"Image.ExportImages.{image_index}_{aw.FileFormatUtil.image_type_to_extension(shape.image_data.image_type)}"
            image_file_path = os.path.join(output_dir, image_file_name)
            shape.image_data.save(image_file_path)
            extracted_images.append(image_file_path)
            image_index += 1

    return extracted_images


def extract_images_from_docx(file_path, output_dir):
    text = docx2txt.process(file_path, output_dir)

    extracted_images = []
    image_index = 0

    for image_file in os.listdir(output_dir):
        if image_file.endswith(".png") or image_file.endswith(".jpg"):
            image_file_path = os.path.join(output_dir, image_file)
            extracted_images.append(image_file_path)
            image_index += 1

    return extracted_images


def perform_ocr(image_file_path):
    reader = easyocr.Reader(['en'])
    result = reader.readtext(image_file_path, detail=0)
    return result


def extract_images_and_perform_ocr(folder_path):

    extracted_images_dir = "./extracted_images"
    if not os.path.exists(extracted_images_dir):
        os.makedirs(extracted_images_dir)

    all_extracted_images = []
    all_ocr_results = []

    # Extract images from PDF files and Document files
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith('.pdf'):
            extracted_images = extract_images_from_pdf(
                file_path, extracted_images_dir)
            all_extracted_images.extend(extracted_images)
        elif filename.endswith('.docx'):
            extracted_images = extract_images_from_docx(
                file_path, extracted_images_dir)
            all_extracted_images.extend(extracted_images)

    # Perform OCR on the extracted images
    for image_path in all_extracted_images:
        ocr_result = perform_ocr(image_path)
        all_ocr_results.append(ocr_result)
        # #all_ocr_results.extend(ocr_result)

    ocr_result_text = "\n\n\n".join(str(result) for result in all_ocr_results)
    ingestText(ocr_result_text)
    # Join the OCR results into a single string
    # ocr_result_text = "  ".join(all_ocr_results)
    # print("OCR Results:")
    # print(ocr_result_text)


# ─── Calling The Functions ────────────────────────────────────────────────────

youtube_links = ["https://www.youtube.com/watch?v=Ch6zNEq9fwM",
                 "https://www.youtube.com/watch?v=V4Z8EdiJxgk&t"]
web_links = ["https://en.wikipedia.org/wiki/Peafowl"]

ingestTextFiles(folder_path)
ingestPDF(folder_path)
ingestCSV(folder_path)
ingestExcel(folder_path)
ingest_youtube_video(youtube_links)
ingestWebpage(web_links)
ingestDocx(folder_path)
extract_images_and_perform_ocr(folder_path)
